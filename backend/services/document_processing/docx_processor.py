"""
DOCX Processor

Extract text and structure from Word documents (.docx format).
"""

import logging
from typing import Dict, Any, List, Optional
from pathlib import Path
import re

try:
    from docx import Document
    from docx.document import Document as DocumentType
    from docx.text.paragraph import Paragraph
    from docx.table import Table
except ImportError:
    raise ImportError(
        "python-docx is required for DOCX processing. "
        "Install with: pip install python-docx"
    )

logger = logging.getLogger(__name__)


class DOCXProcessor:
    """
    Process DOCX (Word) documents to extract text and structure.

    Features:
    - Extract full text content
    - Extract metadata (title, author, date)
    - Parse document structure (headings, paragraphs)
    - Extract tables
    - Identify legislative structure (articles, paragraphs, points)
    """

    def process_docx_from_file(self, file_path: Path) -> Dict[str, Any]:
        """
        Process DOCX file from filesystem.

        Args:
            file_path: Path to DOCX file

        Returns:
            Dict with:
            - text: Full document text
            - metadata: Document metadata
            - structure: Document structure (sections, headings)
            - tables: Extracted tables
            - quality: Processing quality assessment

        Example:
            >>> processor = DOCXProcessor()
            >>> content = processor.process_docx_from_file(Path("document.docx"))
            >>> print(content['text'])
        """
        try:
            logger.info(f"Processing DOCX file: {file_path}")

            # Load document
            doc = Document(str(file_path))

            # Extract components
            metadata = self._extract_metadata(doc)
            text = self._extract_text(doc)
            structure = self._extract_structure(doc)
            tables = self._extract_tables(doc)

            # Assess quality
            quality = self._assess_quality(text, structure)

            result = {
                'text': text,
                'metadata': metadata,
                'structure': structure,
                'tables': tables,
                'quality': quality,
                'processor': 'docx',
                'file_path': str(file_path)
            }

            logger.info(
                f"Successfully processed DOCX: {len(text)} chars, "
                f"{len(structure.get('sections', []))} sections, "
                f"{len(tables)} tables"
            )

            return result

        except Exception as e:
            logger.error(f"Error processing DOCX file {file_path}: {str(e)}")
            raise

    def process_docx_from_bytes(self, file_bytes: bytes, filename: str = "") -> Dict[str, Any]:
        """
        Process DOCX from bytes (uploaded file).

        Args:
            file_bytes: DOCX file content as bytes
            filename: Original filename (for logging)

        Returns:
            Same format as process_docx_from_file
        """
        try:
            logger.info(f"Processing DOCX from bytes: {filename}")

            # Load document from bytes
            from io import BytesIO
            doc = Document(BytesIO(file_bytes))

            # Extract components
            metadata = self._extract_metadata(doc)
            text = self._extract_text(doc)
            structure = self._extract_structure(doc)
            tables = self._extract_tables(doc)

            # Assess quality
            quality = self._assess_quality(text, structure)

            result = {
                'text': text,
                'metadata': metadata,
                'structure': structure,
                'tables': tables,
                'quality': quality,
                'processor': 'docx',
                'filename': filename
            }

            logger.info(
                f"Successfully processed DOCX: {len(text)} chars, "
                f"{len(structure.get('sections', []))} sections"
            )

            return result

        except Exception as e:
            logger.error(f"Error processing DOCX from bytes: {str(e)}")
            raise

    def _extract_metadata(self, doc: DocumentType) -> Dict[str, Any]:
        """Extract document metadata"""
        try:
            core_props = doc.core_properties

            metadata = {
                'title': core_props.title or '',
                'author': core_props.author or '',
                'subject': core_props.subject or '',
                'keywords': core_props.keywords or '',
                'created': core_props.created.isoformat() if core_props.created else None,
                'modified': core_props.modified.isoformat() if core_props.modified else None,
                'category': core_props.category or '',
                'comments': core_props.comments or ''
            }

            return metadata

        except Exception as e:
            logger.warning(f"Error extracting metadata: {str(e)}")
            return {}

    def _extract_text(self, doc: DocumentType) -> str:
        """Extract all text from document, skipping preamble"""
        try:
            text_parts = []

            # Find where legislative content starts
            start_index = self._detect_legislative_start(doc.paragraphs)

            # Extract text from legislative start onwards
            for i, para in enumerate(doc.paragraphs):
                if i >= start_index and para.text.strip():
                    text_parts.append(para.text)

            return '\n\n'.join(text_parts)

        except Exception as e:
            logger.error(f"Error extracting text: {str(e)}")
            return ""

    def _extract_structure(self, doc: DocumentType) -> Dict[str, Any]:
        """
        Extract document structure (headings, sections).

        Returns:
            Dict with:
            - sections: List of sections with headings and content
            - legislative_structure: Detected articles, recitals, etc.
        """
        try:
            sections = []
            current_section = None

            for para in doc.paragraphs:
                # Check if it's a heading
                if para.style.name.startswith('Heading'):
                    # Save previous section
                    if current_section:
                        sections.append(current_section)

                    # Start new section
                    level = self._get_heading_level(para.style.name)
                    current_section = {
                        'type': 'heading',
                        'level': level,
                        'text': para.text,
                        'content': []
                    }

                else:
                    # Add to current section
                    if current_section:
                        current_section['content'].append(para.text)
                    else:
                        # No heading yet, create default section
                        if not sections:
                            current_section = {
                                'type': 'section',
                                'level': 0,
                                'text': 'Introduction',
                                'content': [para.text]
                            }

            # Add last section
            if current_section:
                sections.append(current_section)

            # Detect legislative structure
            legislative_structure = self._detect_legislative_structure(doc)

            return {
                'sections': sections,
                'legislative_structure': legislative_structure,
                'total_sections': len(sections)
            }

        except Exception as e:
            logger.error(f"Error extracting structure: {str(e)}")
            return {'sections': [], 'legislative_structure': {}}

    def _detect_legislative_start(self, paragraphs: List[Paragraph]) -> int:
        """
        Detect where legislative content begins (first recital or article).

        This skips the preamble (Having regard to..., Whereas..., etc.)
        and starts from the first numbered recital or article.

        Args:
            paragraphs: List of document paragraphs

        Returns:
            Index of first legislative paragraph (0 if not found)
        """
        recital_pattern = re.compile(r'^\(1\)\s+', re.IGNORECASE)
        article_pattern = re.compile(r'^Article\s+1\b', re.IGNORECASE)

        for i, para in enumerate(paragraphs):
            text = para.text.strip()

            # Check if this is the first recital
            if recital_pattern.match(text):
                logger.info(f"Legislative content starts at paragraph {i} (Recital 1)")
                return i

            # Check if this is Article 1 (for documents without recitals)
            if article_pattern.match(text):
                logger.info(f"Legislative content starts at paragraph {i} (Article 1)")
                return i

        # If not found, start from beginning (fail-safe)
        logger.warning("Could not detect legislative start, using entire document")
        return 0

    def _detect_legislative_structure(self, doc: DocumentType) -> Dict[str, List[Dict[str, Any]]]:
        """
        Detect legislative elements (articles, recitals, points, paragraphs, subparagraphs)

        This creates a comprehensive structure for AT4AM-style amendment editing.
        Each element becomes one row in the two-column amendment table.

        Returns:
            Dict with:
            - elements: List of all legislative elements in order
            - recitals: List of recitals
            - articles: List of articles with nested structure
            - has_legislative_structure: Boolean
        """
        elements = []
        recitals = []
        articles = []

        # Patterns for legislative elements
        recital_pattern = re.compile(r'^\((\d+)\)\s*(.*)', re.DOTALL)
        article_num_pattern = re.compile(r'^Article\s+(\d+[a-z]?)\s*$', re.IGNORECASE)
        article_title_pattern = re.compile(r'^[A-Z][A-Za-z\s,\-\']+$')  # Article title (Title Case or ALL CAPS)
        point_pattern = re.compile(r'^(\d+)\.\s+(.*)', re.DOTALL)
        paragraph_pattern = re.compile(r'^\(([a-z])\)\s+(.*)', re.DOTALL)
        subparagraph_pattern = re.compile(r'^\(([ivxl]+)\)\s+(.*)', re.DOTALL)
        chapter_pattern = re.compile(r'^(Chapter|CHAPTER|Title|TITLE)\s+([IVX]+|\d+)', re.IGNORECASE)

        current_recital = None
        current_article = None
        current_point = None
        current_paragraph = None
        current_element_text = []

        paragraphs = doc.paragraphs

        # Find where legislative content starts
        start_index = self._detect_legislative_start(paragraphs)

        # Process only from legislative start onwards
        for i, para in enumerate(paragraphs):
            if i < start_index:
                continue

            text = para.text.strip()

            if not text:
                continue

            # Check for Chapter/Title
            chapter_match = chapter_pattern.match(text)
            if chapter_match:
                element = {
                    'type': 'chapter',
                    'number': chapter_match.group(2),
                    'text': text,
                    'level': 0
                }
                elements.append(element)
                continue

            # Check for Recital
            recital_match = recital_pattern.match(text)
            if recital_match:
                # Save previous recital if exists
                if current_recital:
                    current_recital['text'] = ' '.join(current_element_text).strip()
                    recitals.append(current_recital)
                    elements.append(current_recital)

                # Start new recital
                current_recital = {
                    'type': 'recital',
                    'number': recital_match.group(1),
                    'text': recital_match.group(2).strip(),
                    'level': 0
                }
                current_element_text = [recital_match.group(2).strip()]
                current_article = None
                current_point = None
                current_paragraph = None
                continue

            # Check for Article number
            article_num_match = article_num_pattern.match(text)
            if article_num_match:
                # Save previous recital if exists
                if current_recital:
                    current_recital['text'] = ' '.join(current_element_text).strip()
                    recitals.append(current_recital)
                    elements.append(current_recital)
                    current_recital = None

                # Start new article
                current_article = {
                    'type': 'article',
                    'number': article_num_match.group(1),
                    'text': text,
                    'title': '',
                    'content': [],
                    'level': 0
                }
                articles.append(current_article)
                elements.append(current_article)
                current_point = None
                current_paragraph = None
                current_element_text = []
                continue

            # Check for Article title (next line after Article number)
            if current_article and not current_article.get('title') and article_title_pattern.match(text):
                current_article['title'] = text
                title_element = {
                    'type': 'article_title',
                    'article_number': current_article['number'],
                    'text': text,
                    'level': 1
                }
                elements.append(title_element)
                continue

            # Check for Point (1., 2., 3.)
            point_match = point_pattern.match(text)
            if point_match and current_article:
                # Save previous recital if exists
                if current_recital:
                    current_recital['text'] = ' '.join(current_element_text).strip()
                    recitals.append(current_recital)
                    elements.append(current_recital)
                    current_recital = None

                point_element = {
                    'type': 'point',
                    'number': point_match.group(1),
                    'text': point_match.group(2).strip(),
                    'article_number': current_article['number'],
                    'level': 1
                }
                current_article['content'].append(point_element)
                elements.append(point_element)
                current_point = point_element
                current_paragraph = None
                continue

            # Check for Paragraph ((a), (b), (c))
            paragraph_match = paragraph_pattern.match(text)
            if paragraph_match:
                # Save previous recital if exists
                if current_recital:
                    current_recital['text'] = ' '.join(current_element_text).strip()
                    recitals.append(current_recital)
                    elements.append(current_recital)
                    current_recital = None

                paragraph_element = {
                    'type': 'paragraph',
                    'letter': paragraph_match.group(1),
                    'text': paragraph_match.group(2).strip(),
                    'level': 2 if current_point else 1
                }

                if current_point:
                    paragraph_element['point_number'] = current_point['number']
                    paragraph_element['article_number'] = current_article['number'] if current_article else None

                elements.append(paragraph_element)
                current_paragraph = paragraph_element
                continue

            # Check for Subparagraph ((i), (ii), (iii))
            subparagraph_match = subparagraph_pattern.match(text)
            if subparagraph_match:
                # Save previous recital if exists
                if current_recital:
                    current_recital['text'] = ' '.join(current_element_text).strip()
                    recitals.append(current_recital)
                    elements.append(current_recital)
                    current_recital = None

                subparagraph_element = {
                    'type': 'subparagraph',
                    'roman': subparagraph_match.group(1),
                    'text': subparagraph_match.group(2).strip(),
                    'level': 3
                }

                if current_paragraph:
                    subparagraph_element['paragraph_letter'] = current_paragraph['letter']
                if current_point:
                    subparagraph_element['point_number'] = current_point['number']
                if current_article:
                    subparagraph_element['article_number'] = current_article['number']

                elements.append(subparagraph_element)
                continue

            # If we're in a recital and this doesn't match any pattern, it's continuation
            if current_recital:
                current_element_text.append(text)

        # Save last recital if exists
        if current_recital:
            current_recital['text'] = ' '.join(current_element_text).strip()
            recitals.append(current_recital)
            elements.append(current_recital)

        return {
            'elements': elements,
            'recitals': recitals,
            'articles': articles,
            'total_elements': len(elements),
            'has_legislative_structure': len(elements) > 0
        }

    def _get_heading_level(self, style_name: str) -> int:
        """Extract heading level from style name"""
        try:
            # Extract number from 'Heading 1', 'Heading 2', etc.
            match = re.search(r'\d+', style_name)
            if match:
                return int(match.group())
            return 1
        except Exception:
            return 1

    def _extract_tables(self, doc: DocumentType) -> List[Dict[str, Any]]:
        """Extract tables from document"""
        try:
            tables_data = []

            for table_idx, table in enumerate(doc.tables):
                table_data = {
                    'index': table_idx,
                    'rows': len(table.rows),
                    'columns': len(table.columns),
                    'data': []
                }

                # Extract table data
                for row in table.rows:
                    row_data = []
                    for cell in row.cells:
                        row_data.append(cell.text.strip())
                    table_data['data'].append(row_data)

                tables_data.append(table_data)

            return tables_data

        except Exception as e:
            logger.error(f"Error extracting tables: {str(e)}")
            return []

    def _assess_quality(self, text: str, structure: Dict[str, Any]) -> str:
        """
        Assess extraction quality.

        Returns:
            'high', 'medium', or 'low'
        """
        try:
            # Check text length
            if len(text) < 100:
                return 'low'

            # Check structure
            num_sections = len(structure.get('sections', []))
            has_legislative = structure.get('legislative_structure', {}).get('has_legislative_structure', False)

            if num_sections > 0 or has_legislative:
                return 'high'
            elif len(text) > 1000:
                return 'medium'
            else:
                return 'low'

        except Exception:
            return 'unknown'


# Global singleton
_docx_processor: Optional[DOCXProcessor] = None


def get_docx_processor() -> DOCXProcessor:
    """Get global DOCX processor instance"""
    global _docx_processor

    if _docx_processor is None:
        _docx_processor = DOCXProcessor()

    return _docx_processor
