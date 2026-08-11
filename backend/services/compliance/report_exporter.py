"""
Report Exporter Service

Generates professional DOCX compliance reports with Brubru branding.
"""

import logging
from typing import Optional
from pathlib import Path
from datetime import datetime
from sqlalchemy import case
from sqlalchemy.orm import Session

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

from models.compliance import ComplianceAnalysis, GapFinding
from models.eu_law import LawRequirement, LawCluster, EULaw

logger = logging.getLogger(__name__)


class ReportExporter:
    """
    Generates DOCX compliance reports with Brubru branding.
    
    Features:
    - Professional formatting with Brubru colors
    - Compliance score summary
    - Gap findings organized by priority
    - Action plan with deadlines
    - Watermark for Yellow tier users
    """
    
    # Brubru brand colors
    BRUBRU_PURPLE = RGBColor(102, 126, 234)  # #667eea
    BRUBRU_DARK = RGBColor(45, 55, 72)       # #2d3748
    
    STATUS_COLORS = {
        'met': RGBColor(72, 187, 120),       # Green #48bb78
        'partial': RGBColor(221, 107, 32),   # Orange #dd6b20
        'gap': RGBColor(197, 48, 48),        # Red #c53030
    }
    
    def __init__(self, db: Session):
        if not DOCX_AVAILABLE:
            raise RuntimeError("python-docx not installed. Install with: pip install python-docx")
        
        self.db = db
    
    def export_analysis(
        self,
        analysis_id: int,
        output_path: str,
        include_watermark: bool = False
    ) -> str:
        """
        Export compliance analysis to DOCX.
        
        Args:
            analysis_id: ID of the compliance analysis
            output_path: Path to save the DOCX file
            include_watermark: Add watermark for Yellow tier users
            
        Returns:
            Path to the generated file
        """
        logger.info(f"Generating DOCX report for analysis {analysis_id}")

        analysis, cluster, findings = self._load(analysis_id)

        # Create document
        doc = Document()
        
        # Add watermark if needed
        if include_watermark:
            self._add_watermark(doc)
        
        # Build report
        self._add_header(doc, cluster, analysis)
        self._add_executive_summary(doc, analysis)
        self._add_compliance_score(doc, analysis)
        self._add_gap_findings(doc, findings)
        self._add_action_plan(doc, findings)
        self._add_footer(doc)
        
        # Save document
        doc.save(output_path)
        logger.info(f"Report saved to {output_path}")
        
        return output_path
    

    # ------------------------------------------------------------------
    # PDF
    # ------------------------------------------------------------------

    def export_analysis_pdf(
        self,
        analysis_id: int,
        output_path: str,
        include_watermark: bool = False
    ) -> str:
        """Export the same compliance analysis to PDF.

        POST /analysis/{id}/export?export_format=pdf returned 501 until
        8 Aug 2026. Built on reportlab, which was already a dependency for
        user_documents export (and is now also in requirements-light.txt, which
        is what production actually installs -- it was only in requirements.txt
        before, so this would have raised ImportError on Railway).

        Content is loaded through the same _load() as the DOCX path so the two
        formats cannot drift apart.
        """
        from reportlab.lib import colors
        from reportlab.lib.enums import TA_CENTER
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import cm
        from reportlab.platypus import (
            SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak,
        )
        from xml.sax.saxutils import escape

        logger.info(f"Generating PDF report for analysis {analysis_id}")
        analysis, cluster, findings = self._load(analysis_id)

        purple = colors.HexColor("#667eea")
        dark = colors.HexColor("#2d3748")
        status_colour = {
            "met": colors.HexColor("#2e9e6b"),
            "partial": colors.HexColor("#b5751a"),
            "gap": colors.HexColor("#c53030"),
            "not_applicable": colors.HexColor("#6b7280"),
        }

        styles = getSampleStyleSheet()
        h1 = ParagraphStyle("BH1", parent=styles["Heading1"], textColor=purple, spaceAfter=10)
        h2 = ParagraphStyle("BH2", parent=styles["Heading2"], textColor=dark, spaceBefore=12, spaceAfter=6)
        body = ParagraphStyle("BBody", parent=styles["BodyText"], fontSize=9.5, leading=14)
        small = ParagraphStyle("BSmall", parent=styles["BodyText"], fontSize=8, leading=11,
                               textColor=colors.HexColor("#55607a"))
        score_style = ParagraphStyle("BScore", parent=styles["Title"], fontSize=42, alignment=TA_CENTER,
                                     textColor=purple, spaceAfter=2)

        def on_page(canvas, doc_):
            canvas.saveState()
            if include_watermark:
                canvas.setFont("Helvetica-Bold", 54)
                canvas.setFillColor(colors.HexColor("#e8e8e8"))
                canvas.translate(A4[0] / 2, A4[1] / 2)
                canvas.rotate(45)
                canvas.drawCentredString(0, 0, "BRUBRU")
                canvas.rotate(-45)
                canvas.translate(-A4[0] / 2, -A4[1] / 2)
            canvas.setFont("Helvetica", 7.5)
            canvas.setFillColor(colors.HexColor("#8a93a3"))
            canvas.drawString(2 * cm, 1.2 * cm, "Generated by Brubru - EU Law Comply")
            canvas.drawRightString(A4[0] - 2 * cm, 1.2 * cm, f"Page {doc_.page}")
            canvas.restoreState()

        doc = SimpleDocTemplate(
            output_path, pagesize=A4,
            leftMargin=2 * cm, rightMargin=2 * cm, topMargin=2 * cm, bottomMargin=2 * cm,
            title=f"Brubru EU Compliance Report - {cluster.name if cluster else analysis_id}",
            author="Brubru",
        )

        story = []
        story.append(Paragraph("Brubru - EU Law Comply", h1))
        story.append(Paragraph("Your EU Law Compliance Report", h2))
        story.append(Spacer(1, 0.4 * cm))

        meta = [
            ["Analysis", escape(analysis.analysis_name or "Unnamed analysis")],
            ["Law cluster", escape(cluster.name if cluster else "-")],
            ["Policy area", escape((cluster.policy_area if cluster else None) or "-")],
            ["Report date", datetime.utcnow().strftime("%Y-%m-%d %H:%M UTC")],
        ]
        t = Table(meta, colWidths=[4 * cm, 11 * cm])
        t.setStyle(TableStyle([
            ("FONTSIZE", (0, 0), (-1, -1), 9),
            ("TEXTCOLOR", (0, 0), (0, -1), dark),
            ("FONTNAME", (0, 0), (0, -1), "Helvetica-Bold"),
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
            ("LINEBELOW", (0, 0), (-1, -2), 0.4, colors.HexColor("#e6eaf0")),
        ]))
        story.append(t)
        story.append(Spacer(1, 0.7 * cm))

        # Score. An analysis with no score is a failed run, not zero
        # compliance -- say so rather than printing 0.0%.
        story.append(Paragraph("Compliance score", h2))
        if analysis.compliance_score is None:
            story.append(Paragraph(
                "This analysis did not complete, so no compliance score was produced. "
                "This is not a score of zero.", body))
        else:
            story.append(Paragraph(f"{float(analysis.compliance_score):.1f}%", score_style))
            story.append(Paragraph("Overall compliance", small))
            story.append(Spacer(1, 0.3 * cm))
            counts = Table([
                ["Requirements met", "Partial", "Gaps", "Total analysed"],
                [str(analysis.requirements_met or 0), str(analysis.requirements_partial or 0),
                 str(analysis.requirements_gap or 0), str(analysis.total_requirements or 0)],
            ], colWidths=[3.75 * cm] * 4)
            counts.setStyle(TableStyle([
                ("FONTSIZE", (0, 0), (-1, -1), 9),
                ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                ("ALIGN", (0, 0), (-1, -1), "CENTER"),
                ("TEXTCOLOR", (0, 1), (0, 1), status_colour["met"]),
                ("TEXTCOLOR", (1, 1), (1, 1), status_colour["partial"]),
                ("TEXTCOLOR", (2, 1), (2, 1), status_colour["gap"]),
                ("FONTSIZE", (0, 1), (-1, 1), 14),
                ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#e6eaf0")),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
                ("TOPPADDING", (0, 0), (-1, -1), 6),
            ]))
            story.append(counts)

        if findings:
            story.append(PageBreak())
            story.append(Paragraph("Findings", h2))
            story.append(Paragraph(
                "Ordered by priority, then by severity. Evidence is quoted from the documents "
                "you uploaded; where nothing addressed a requirement, that is stated explicitly.",
                small))
            story.append(Spacer(1, 0.35 * cm))

            for finding, requirement, law in findings:
                sc = status_colour.get(finding.status, dark)
                head = ParagraphStyle("fh", parent=body, fontSize=10, textColor=sc,
                                      fontName="Helvetica-Bold", spaceBefore=8)
                story.append(Paragraph(
                    f"{escape(str(finding.status).replace('_', ' ')).upper()} &nbsp;|&nbsp; "
                    f"{escape(requirement.article or '-')} &nbsp;|&nbsp; "
                    f"{escape(requirement.criticality or '-')}", head))
                story.append(Paragraph(escape(requirement.requirement_text or ""), body))
                if law is not None and law.celex:
                    story.append(Paragraph(f"Source: {escape(law.celex)}", small))
                if finding.evidence_text:
                    story.append(Paragraph(
                        f"<b>Evidence:</b> &ldquo;{escape(finding.evidence_text)}&rdquo;"
                        + (f" &mdash; {escape(finding.evidence_source)}" if finding.evidence_source else ""),
                        small))
                else:
                    story.append(Paragraph(
                        "<b>Evidence:</b> nothing in the uploaded documents addresses this requirement.",
                        small))
                if finding.gap_description:
                    story.append(Paragraph(f"<b>Gap:</b> {escape(finding.gap_description)}", small))
                if finding.recommendation:
                    story.append(Paragraph(f"<b>Recommendation:</b> {escape(finding.recommendation)}", small))
                if finding.confidence_score is not None:
                    story.append(Paragraph(
                        f"<b>Confidence:</b> {float(finding.confidence_score):.0f}%", small))
                story.append(Spacer(1, 0.15 * cm))

        doc.build(story, onFirstPage=on_page, onLaterPages=on_page)
        logger.info(f"PDF report saved to {output_path}")
        return output_path

    def _load(self, analysis_id: int):
        """Load the analysis, its cluster and its findings.

        Shared by the DOCX and PDF exporters so the two formats can never
        disagree about ordering or content.
        """
        analysis = self.db.query(ComplianceAnalysis).filter(
            ComplianceAnalysis.id == analysis_id
        ).first()
        if not analysis:
            raise ValueError(f"Analysis {analysis_id} not found")

        cluster = self.db.query(LawCluster).filter(
            LawCluster.id == analysis.cluster_id
        ).first()

        findings = self.db.query(GapFinding, LawRequirement, EULaw).join(
            LawRequirement, GapFinding.requirement_id == LawRequirement.id
        ).join(
            EULaw, LawRequirement.law_id == EULaw.id
        ).filter(
            GapFinding.analysis_id == analysis_id
        ).order_by(
            GapFinding.priority,
            # criticality is free text: DESC sorts alphabetically and puts
            # 'recommended' above 'critical'. Order by severity explicitly.
            # Mirrors CRITICALITY_ORDER in api/eu_law_comply.py.
            case(
                (LawRequirement.criticality == 'critical', 0),
                (LawRequirement.criticality == 'important', 1),
                (LawRequirement.criticality == 'recommended', 2),
                else_=3,
            )
        ).all()
        return analysis, cluster, findings

    def _add_watermark(self, doc: Document):
        """Add 'BRUBRU - YELLOW TIER' watermark."""
        section = doc.sections[0]
        header = section.header
        
        paragraph = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        run = paragraph.add_run('BRUBRU - YELLOW TIER')
        run.font.size = Pt(48)
        run.font.color.rgb = RGBColor(200, 200, 200)
        run.font.bold = True
        
        # Make semi-transparent (requires additional XML manipulation)
        # This is a simplified version - full watermark requires more complex XML
    
    def _add_header(self, doc: Document, cluster: LawCluster, analysis: ComplianceAnalysis):
        """Add report header with title and metadata."""
        # Title
        title = doc.add_heading('Brubru - EU Law Comply', level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title.runs[0]
        title_run.font.color.rgb = self.BRUBRU_PURPLE
        
        # Subtitle
        subtitle = doc.add_heading('Your EU Law Compliance Report', level=2)
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle_run = subtitle.runs[0]
        subtitle_run.font.color.rgb = self.BRUBRU_DARK
        
        # Metadata
        doc.add_paragraph()
        meta_table = doc.add_table(rows=4, cols=2)
        meta_table.style = 'Light Grid Accent 1'
        
        meta_table.rows[0].cells[0].text = 'Analysis Name:'
        meta_table.rows[0].cells[1].text = analysis.analysis_name or 'Unnamed Analysis'
        
        meta_table.rows[1].cells[0].text = 'Law Cluster:'
        meta_table.rows[1].cells[1].text = cluster.name
        
        meta_table.rows[2].cells[0].text = 'Policy Area:'
        meta_table.rows[2].cells[1].text = cluster.policy_area or 'N/A'
        
        meta_table.rows[3].cells[0].text = 'Report Date:'
        meta_table.rows[3].cells[1].text = datetime.utcnow().strftime('%Y-%m-%d %H:%M UTC')
        
        doc.add_page_break()
    
    def _add_executive_summary(self, doc: Document, analysis: ComplianceAnalysis):
        """Add executive summary section."""
        doc.add_heading('Executive Summary', level=1)
        
        summary_text = f"""This report provides a comprehensive compliance analysis against the {analysis.cluster_id} 
regulatory package. A total of {analysis.total_requirements or 0} requirements were analyzed 
against your organization's documentation."""
        
        doc.add_paragraph(summary_text)
        doc.add_paragraph()
    
    def _add_compliance_score(self, doc: Document, analysis: ComplianceAnalysis):
        """Add overall compliance score with breakdown."""
        doc.add_heading('Compliance Score', level=1)
        
        # Overall score
        score_para = doc.add_paragraph()
        score_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        score_run = score_para.add_run(f'{analysis.compliance_score:.1f}%' if analysis.compliance_score else 'N/A')
        score_run.font.size = Pt(48)
        score_run.font.bold = True
        
        # Determine color based on score
        if analysis.compliance_score:
            if analysis.compliance_score >= 90:
                score_run.font.color.rgb = self.STATUS_COLORS['met']
            elif analysis.compliance_score >= 70:
                score_run.font.color.rgb = self.STATUS_COLORS['partial']
            else:
                score_run.font.color.rgb = self.STATUS_COLORS['gap']
        
        doc.add_paragraph()
        
        # Breakdown table
        breakdown_table = doc.add_table(rows=4, cols=2)
        breakdown_table.style = 'Light Grid Accent 1'
        
        breakdown_table.rows[0].cells[0].text = '✓ Requirements Met'
        breakdown_table.rows[0].cells[1].text = str(analysis.requirements_met or 0)
        
        breakdown_table.rows[1].cells[0].text = '⚠ Partial Compliance'
        breakdown_table.rows[1].cells[1].text = str(analysis.requirements_partial or 0)
        
        breakdown_table.rows[2].cells[0].text = '✗ Critical Gaps'
        breakdown_table.rows[2].cells[1].text = str(analysis.requirements_gap or 0)
        
        breakdown_table.rows[3].cells[0].text = 'Total Requirements'
        breakdown_table.rows[3].cells[1].text = str(analysis.total_requirements or 0)
        
        doc.add_page_break()
    
    def _add_gap_findings(self, doc: Document, findings):
        """Add detailed gap findings section."""
        doc.add_heading('Gap Findings', level=1)
        
        # Group by status.
        #
        # This read `[f for f, req, law in findings if f[0].status == 'gap']`,
        # which is wrong twice: the row is already unpacked so `f` is a
        # GapFinding and `f[0]` raises TypeError, and even past that it yields a
        # bare finding which the loops below unpack as a 3-tuple. DOCX export
        # therefore raised on every call since it was written -- which, together
        # with PDF returning 501, is the whole reason analysis_exports had zero
        # rows. Keep the full tuples.
        gaps = [(f, req, law) for f, req, law in findings if f.status == 'gap']
        partials = [(f, req, law) for f, req, law in findings if f.status == 'partial']
        
        # Critical gaps
        if gaps:
            doc.add_heading('Critical Gaps', level=2)
            for finding, requirement, law in gaps:
                self._add_finding_detail(doc, finding, requirement, law)
        
        # Partial compliance
        if partials:
            doc.add_heading('Partial Compliance', level=2)
            for finding, requirement, law in partials:
                self._add_finding_detail(doc, finding, requirement, law)
        
        if not gaps and not partials:
            doc.add_paragraph('✓ Excellent! No critical gaps or partial compliance issues identified.')
    
    def _add_finding_detail(self, doc: Document, finding, requirement, law):
        """Add details for a single finding."""
        # Article and requirement
        heading = doc.add_heading(f'{requirement.article} - {law.title[:100]}', level=3)
        
        # Priority badge
        priority_para = doc.add_paragraph()
        priority_run = priority_para.add_run(f'Priority: {finding.priority}/5')
        priority_run.bold = True
        
        if finding.priority <= 2:
            priority_run.font.color.rgb = self.STATUS_COLORS['gap']
        elif finding.priority <= 3:
            priority_run.font.color.rgb = self.STATUS_COLORS['partial']
        
        # Requirement text
        doc.add_paragraph(f'Requirement: {requirement.requirement_text}')
        
        # Gap description
        if finding.gap_description:
            gap_para = doc.add_paragraph()
            gap_run = gap_para.add_run(f'Gap: {finding.gap_description}')
            gap_run.italic = True
        
        # Recommendation
        if finding.recommendation:
            rec_para = doc.add_paragraph()
            rec_run = rec_para.add_run(f'Recommendation: {finding.recommendation}')
            rec_run.bold = True
        
        # Effort estimate
        if finding.estimated_effort:
            doc.add_paragraph(f'Estimated Effort: {finding.estimated_effort.title()}')
        
        doc.add_paragraph()  # Spacing
    
    def _add_action_plan(self, doc: Document, findings):
        """Add action plan timeline."""
        doc.add_page_break()
        doc.add_heading('Action Plan', level=1)
        
        # Filter gaps and partials, sort by priority
        action_items = [
            (finding, requirement, law)
            for finding, requirement, law in findings
            if finding.status in ['gap', 'partial']
        ]
        
        if not action_items:
            doc.add_paragraph('✓ No action items required. All requirements are met!')
            return
        
        # Group by priority
        for priority in range(1, 6):
            priority_items = [
                (f, r, l) for f, r, l in action_items if f.priority == priority
            ]
            
            if priority_items:
                priority_labels = {1: 'Critical', 2: 'High', 3: 'Medium', 4: 'Low', 5: 'Optional'}
                doc.add_heading(f'{priority_labels[priority]} Priority Actions', level=2)
                
                for finding, requirement, law in priority_items:
                    # Action item
                    para = doc.add_paragraph(style='List Bullet')
                    para.add_run(f'{requirement.article}: ').bold = True
                    para.add_run(finding.recommendation or 'Address compliance gap')
                    
                    # Deadline if available
                    if requirement.deadline:
                        deadline_para = doc.add_paragraph(style='List Bullet 2')
                        deadline_para.add_run(f'  Deadline: {requirement.deadline.strftime("%Y-%m-%d")}')
    
    def _add_footer(self, doc: Document):
        """Add report footer."""
        doc.add_page_break()
        
        footer_para = doc.add_paragraph()
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        footer_run = footer_para.add_run('Generated by Brubru - EU Law Comply')
        footer_run.italic = True
        footer_run.font.size = Pt(10)
        footer_run.font.color.rgb = RGBColor(100, 100, 100)
        
        date_para = doc.add_paragraph()
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        date_run = date_para.add_run(datetime.utcnow().strftime('%Y-%m-%d %H:%M UTC'))
        date_run.font.size = Pt(9)
        date_run.font.color.rgb = RGBColor(150, 150, 150)
