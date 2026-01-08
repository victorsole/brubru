"""
Amendment Export Service

Service for exporting amendments to DOCX format with proper table formatting.
"""

import logging
from pathlib import Path
from datetime import datetime
from typing import List
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

logger = logging.getLogger(__name__)


class AmendmentExportService:
    """Service for exporting amendments to DOCX format"""

    def __init__(self, export_dir: str = "backend/cache/exports"):
        """
        Initialize the export service.

        Args:
            export_dir: Directory to store exported files
        """
        self.export_dir = Path(export_dir)
        self.export_dir.mkdir(parents=True, exist_ok=True)

    def export_amendments_to_docx(
        self,
        amendments: List[dict],
        document_celex: str,
        user_id: str
    ) -> Path:
        """
        Export amendments to a DOCX file with table format.

        Args:
            amendments: List of amendment dictionaries
            document_celex: CELEX number of the source document
            user_id: User ID for file naming

        Returns:
            Path to the generated DOCX file
        """
        try:
            # Create document
            doc = Document()

            # Set document properties
            doc.core_properties.title = f"Amendments - {document_celex}"
            doc.core_properties.author = "Brubru Amendator"
            doc.core_properties.created = datetime.now()

            # Add title
            title = doc.add_heading(f'Amendments - {document_celex}', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Add date
            date_para = doc.add_paragraph()
            date_para.add_run(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}").italic = True
            date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

            doc.add_paragraph()  # Spacer

            # Sort amendments by element index
            sorted_amendments = sorted(amendments, key=lambda a: a.get('element_index', 0))

            # Add each amendment
            for idx, amendment in enumerate(sorted_amendments, start=1):
                self._add_amendment_to_document(doc, amendment, idx)

                # Add separator line
                doc.add_paragraph('_' * 60)
                doc.add_paragraph()  # Spacer

            # Generate filename
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            filename = f"{document_celex}_amendments_{user_id}_{timestamp}.docx"

            # Create user-specific directory
            user_dir = self.export_dir / user_id
            user_dir.mkdir(parents=True, exist_ok=True)

            filepath = user_dir / filename

            # Save document
            doc.save(str(filepath))

            logger.info(f"Exported {len(amendments)} amendments to {filepath}")

            return filepath

        except Exception as e:
            logger.error(f"Error exporting amendments to DOCX: {str(e)}")
            raise

    def _add_amendment_to_document(self, doc: Document, amendment: dict, number: int):
        """
        Add a single amendment to the document with table format.

        Args:
            doc: Document object
            amendment: Amendment dictionary
            number: Amendment number
        """
        # Amendment header
        header = doc.add_heading(f'AMENDMENT {number}', level=2)

        # Position and type info
        info = doc.add_paragraph()
        info.add_run('Position: ').bold = True
        info.add_run(amendment.get('position_text', 'Unknown'))

        info = doc.add_paragraph()
        info.add_run('Type: ').bold = True
        type_text = amendment.get('amendment_type', 'Unknown').capitalize()
        info.add_run(type_text)

        # Create table with 2 columns
        table = doc.add_table(rows=3, cols=2)
        table.style = 'Table Grid'

        # Set column widths
        for row in table.rows:
            for idx, cell in enumerate(row.cells):
                if idx == 0:
                    cell.width = Inches(3.0)
                else:
                    cell.width = Inches(3.5)

        # Header row
        header_cells = table.rows[0].cells
        self._set_cell_properties(header_cells[0], 'Original text', bold=True, bg_color='D9E2F3')
        self._set_cell_properties(header_cells[1], 'Proposed text', bold=True, bg_color='D9E2F3')

        # Content row
        content_cells = table.rows[1].cells

        # Original text
        original_text = amendment.get('original_text', '')
        if amendment.get('amendment_type') == 'addition':
            original_text = '[No original text - Addition]'
        self._set_cell_properties(content_cells[0], original_text)

        # Proposed text
        proposed_text = amendment.get('proposed_text', '')
        if amendment.get('amendment_type') == 'suppression':
            proposed_text = '[Suppressed text]'
            # Add strikethrough to original text if suppression
            para = content_cells[0].paragraphs[0]
            for run in para.runs:
                run.font.strike = True
        self._set_cell_properties(content_cells[1], proposed_text)

        # Merge bottom row for justification
        justification_row = table.rows[2]
        justification_cell = justification_row.cells[0].merge(justification_row.cells[1])

        # Add justification
        justification_text = amendment.get('justification', 'No justification provided')
        self._set_cell_properties(
            justification_cell,
            f"Justification:\n{justification_text}",
            bg_color='F2F2F2'
        )

    def _set_cell_properties(self, cell, text: str, bold: bool = False, bg_color: str = None):
        """
        Set cell text and formatting properties.

        Args:
            cell: Table cell object
            text: Text to add
            bold: Whether text should be bold
            bg_color: Background color in hex (e.g., 'D9E2F3')
        """
        # Clear existing content
        cell.text = ''

        # Add text
        paragraph = cell.paragraphs[0]
        run = paragraph.add_run(text)
        run.font.size = Pt(11)

        if bold:
            run.font.bold = True

        # Set background color
        if bg_color:
            shading_elm = OxmlElement('w:shd')
            shading_elm.set(qn('w:fill'), bg_color)
            cell._element.get_or_add_tcPr().append(shading_elm)

        # Set cell padding
        tc = cell._element
        tcPr = tc.get_or_add_tcPr()
        tcMar = OxmlElement('w:tcMar')

        for margin_name in ['top', 'left', 'bottom', 'right']:
            node = OxmlElement(f'w:{margin_name}')
            node.set(qn('w:w'), '100')
            node.set(qn('w:type'), 'dxa')
            tcMar.append(node)

        tcPr.append(tcMar)


# Global service instance
_export_service: AmendmentExportService = None


def get_export_service() -> AmendmentExportService:
    """Get global export service instance"""
    global _export_service

    if _export_service is None:
        _export_service = AmendmentExportService()

    return _export_service
