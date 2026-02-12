"""
Amendment Export Service

Service for exporting amendments to DOCX format with proper table formatting.
"""

import logging
from pathlib import Path
from datetime import datetime
from typing import List, Optional
import difflib
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

logger = logging.getLogger(__name__)


class AmendmentExportService:
    """Service for exporting amendments to DOCX format"""

    def __init__(self, export_dir: str = "backend/cache/exports"):
        """
        Initialize the export service.

        Args:
            export_dir: Directory to store exported files
        """
        self.export_dir = Path(export_dir)
        self.export_dir.mkdir(parents=True, exist_ok=True)

    def export_amendments_to_docx(
        self,
        amendments: List[dict],
        document_celex: str,
        user_id: str,
        subscription_tier: str = "white",
        document_filename: Optional[str] = None
    ) -> Path:
        """
        Export amendments to a DOCX file with table format.

        Args:
            amendments: List of amendment dictionaries
            document_celex: CELEX number of the source document
            user_id: User ID for file naming
            subscription_tier: User's subscription tier (white, yellow, blue)
            document_filename: Original document filename for display

        Returns:
            Path to the generated DOCX file
        """
        try:
            # Create document
            doc = Document()

            # Set default font to Garamond (closest to Adobe Caslon Pro)
            style = doc.styles['Normal']
            font = style.font
            font.name = 'Garamond'
            font.size = Pt(11)

            # Set document properties
            display_name = document_filename or document_celex
            doc.core_properties.title = f"Amendments - {display_name}"
            doc.core_properties.author = "Brubru Amendator"
            doc.core_properties.created = datetime.now()

            # Add watermark (only for White tier)
            if subscription_tier == "white":
                self._add_watermark(doc)

            # Add header with logo and text
            self._add_header_logo(doc)

            # Add title
            title = doc.add_heading(f'Amendments - {display_name}', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            # Set title font
            for run in title.runs:
                run.font.name = 'Garamond'
                run.font.size = Pt(24)

            # Add date
            date_para = doc.add_paragraph()
            date_run = date_para.add_run(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
            date_run.italic = True
            date_run.font.name = 'Garamond'
            date_run.font.size = Pt(11)
            date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

            doc.add_paragraph()  # Spacer

            # Sort amendments by element index
            sorted_amendments = sorted(amendments, key=lambda a: a.get('element_index', 0))

            # Add each amendment
            for idx, amendment in enumerate(sorted_amendments, start=1):
                self._add_amendment_to_document(doc, amendment, idx)

                # Add separator line
                separator = doc.add_paragraph('_' * 60)
                for run in separator.runs:
                    run.font.name = 'Garamond'
                doc.add_paragraph()  # Spacer

            # Generate filename
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            filename = f"{document_celex}_amendments_{user_id}_{timestamp}.docx"

            # Create user-specific directory
            user_dir = self.export_dir / user_id
            user_dir.mkdir(parents=True, exist_ok=True)

            filepath = user_dir / filename

            # Save document
            doc.save(str(filepath))

            logger.info(f"Exported {len(amendments)} amendments to {filepath}")

            return filepath

        except Exception as e:
            logger.error(f"Error exporting amendments to DOCX: {str(e)}")
            raise

    def _add_amendment_to_document(self, doc: Document, amendment: dict, number: int):
        """
        Add a single amendment to the document with table format.

        Args:
            doc: Document object
            amendment: Amendment dictionary
            number: Amendment number
        """
        # Amendment header
        header = doc.add_heading(f'AMENDMENT {number}', level=2)
        for run in header.runs:
            run.font.name = 'Garamond'

        # Position and type info
        info = doc.add_paragraph()
        pos_run = info.add_run('Position: ')
        pos_run.bold = True
        pos_run.font.name = 'Garamond'
        pos_text_run = info.add_run(amendment.get('position_text', 'Unknown'))
        pos_text_run.font.name = 'Garamond'

        info = doc.add_paragraph()
        type_label_run = info.add_run('Type: ')
        type_label_run.bold = True
        type_label_run.font.name = 'Garamond'
        type_text = amendment.get('amendment_type', 'Unknown').capitalize()
        type_text_run = info.add_run(type_text)
        type_text_run.font.name = 'Garamond'

        # Create table with 2 columns
        table = doc.add_table(rows=3, cols=2)
        table.style = 'Table Grid'

        # Set column widths
        for row in table.rows:
            for idx, cell in enumerate(row.cells):
                if idx == 0:
                    cell.width = Inches(3.0)
                else:
                    cell.width = Inches(3.5)

        # Header row
        header_cells = table.rows[0].cells
        self._set_cell_properties(header_cells[0], 'Original text', bold=True, bg_color='D9E2F3')
        self._set_cell_properties(header_cells[1], 'Proposed text', bold=True, bg_color='D9E2F3')

        # Content row
        content_cells = table.rows[1].cells

        # Original text
        original_text = amendment.get('original_text', '')
        if amendment.get('amendment_type') == 'addition':
            original_text = '[No original text - Addition]'
        self._set_cell_properties(content_cells[0], original_text)

        # Proposed text - Bold italic ONLY for changed portions
        proposed_text = amendment.get('proposed_text', '')
        amendment_type = amendment.get('amendment_type', '')

        if amendment_type == 'suppression':
            proposed_text = '[Suppressed text]'
            self._set_cell_properties(content_cells[1], proposed_text, italic=True)
            # Add strikethrough to original text if suppression
            para = content_cells[0].paragraphs[0]
            for run in para.runs:
                run.font.strike = True
        elif amendment_type == 'addition':
            # For additions, all text is new, so all bold italic
            self._set_cell_properties(content_cells[1], proposed_text, bold=True, italic=True)
        else:
            # For modifications, compare and highlight only changed portions
            self._set_proposed_text_with_diff(content_cells[1], original_text, proposed_text)

        # Merge bottom row for justification
        justification_row = table.rows[2]
        justification_cell = justification_row.cells[0].merge(justification_row.cells[1])

        # Add justification
        justification_text = amendment.get('justification', 'No justification provided')
        self._set_cell_properties(
            justification_cell,
            f"Justification:\n{justification_text}",
            bg_color='F2F2F2'
        )

    def _set_proposed_text_with_diff(self, cell, original_text: str, proposed_text: str):
        """
        Set proposed text with bold italic applied only to changed portions.

        Args:
            cell: Table cell object
            original_text: Original text
            proposed_text: Proposed text
        """
        # Clear existing content
        cell.text = ''

        # Get paragraph
        paragraph = cell.paragraphs[0]

        # Split texts into words
        original_words = original_text.split()
        proposed_words = proposed_text.split()

        # Use SequenceMatcher to find differences
        matcher = difflib.SequenceMatcher(None, original_words, proposed_words)

        # Track position in proposed text
        first_segment = True
        for tag, i1, i2, j1, j2 in matcher.get_opcodes():
            if tag == 'equal':
                # Unchanged text - normal formatting
                text = ' '.join(proposed_words[j1:j2])
                if not first_segment:  # Add space before if not first segment
                    text = ' ' + text
                run = paragraph.add_run(text)
                run.font.name = 'Garamond'
                run.font.size = Pt(11)
                first_segment = False
            elif tag == 'replace' or tag == 'insert':
                # Changed or new text - bold italic
                text = ' '.join(proposed_words[j1:j2])
                if not first_segment:  # Add space before if not first segment
                    text = ' ' + text
                run = paragraph.add_run(text)
                run.font.name = 'Garamond'
                run.font.size = Pt(11)
                run.font.bold = True
                run.font.italic = True
                first_segment = False
            # 'delete' tag is ignored since those words aren't in proposed text

        # Set cell padding
        tc = cell._element
        tcPr = tc.get_or_add_tcPr()
        tcMar = OxmlElement('w:tcMar')

        for margin_name in ['top', 'left', 'bottom', 'right']:
            node = OxmlElement(f'w:{margin_name}')
            node.set(qn('w:w'), '100')
            node.set(qn('w:type'), 'dxa')
            tcMar.append(node)

        tcPr.append(tcMar)

    def _add_watermark(self, doc: Document):
        """
        Add watermark to the document (only for White tier users).
        Uses brubru_icon.png as a semi-transparent background image.
        """
        try:
            # Path to watermark image
            watermark_path = Path("frontend/public/assets/brubru_icon.png")

            if not watermark_path.exists():
                logger.warning(f"Watermark image not found at {watermark_path}")
                return

            # Access the document's sectPr element
            section = doc.sections[0]
            sectPr = section._sectPr

            # Create header part if it doesn't exist
            header = section.header
            header_para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()

            # Add watermark image to header with positioning
            run = header_para.add_run()
            run.add_picture(str(watermark_path), width=Inches(3.0))

            # Center the image
            header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Make header semi-transparent by reducing opacity
            # Note: Full watermark support in python-docx is limited
            # This creates a centered logo in the header

        except Exception as e:
            logger.warning(f"Failed to add watermark: {str(e)}")

    def _add_header_logo(self, doc: Document):
        """
        Add Brubru logo and title text at the beginning of the document.
        """
        try:
            # Path to logo image
            logo_path = Path("frontend/public/assets/brubru_icon_colours.png")

            if not logo_path.exists():
                logger.warning(f"Logo image not found at {logo_path}")
                return

            # Add logo paragraph
            logo_para = doc.add_paragraph()
            logo_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Add logo image
            run = logo_para.add_run()
            run.add_picture(str(logo_path), width=Inches(1.5))

            # Add title text
            title_para = doc.add_paragraph()
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title_run = title_para.add_run("Your Amendments With Brubru")
            title_run.font.name = 'Garamond'
            title_run.font.size = Pt(18)
            title_run.font.bold = True
            title_run.font.color.rgb = RGBColor(6, 147, 227)  # Brubru blue color

            # Add spacer
            doc.add_paragraph()

        except Exception as e:
            logger.warning(f"Failed to add header logo: {str(e)}")

    def _set_cell_properties(self, cell, text: str, bold: bool = False, italic: bool = False, bg_color: str = None):
        """
        Set cell text and formatting properties.

        Args:
            cell: Table cell object
            text: Text to add
            bold: Whether text should be bold
            italic: Whether text should be italic
            bg_color: Background color in hex (e.g., 'D9E2F3')
        """
        # Clear existing content
        cell.text = ''

        # Add text
        paragraph = cell.paragraphs[0]
        run = paragraph.add_run(text)
        run.font.name = 'Garamond'
        run.font.size = Pt(11)

        if bold:
            run.font.bold = True

        if italic:
            run.font.italic = True

        # Set background color
        if bg_color:
            shading_elm = OxmlElement('w:shd')
            shading_elm.set(qn('w:fill'), bg_color)
            cell._element.get_or_add_tcPr().append(shading_elm)

        # Set cell padding
        tc = cell._element
        tcPr = tc.get_or_add_tcPr()
        tcMar = OxmlElement('w:tcMar')

        for margin_name in ['top', 'left', 'bottom', 'right']:
            node = OxmlElement(f'w:{margin_name}')
            node.set(qn('w:w'), '100')
            node.set(qn('w:type'), 'dxa')
            tcMar.append(node)

        tcPr.append(tcMar)


# Global service instance
_export_service: AmendmentExportService = None


def get_export_service() -> AmendmentExportService:
    """Get global export service instance"""
    global _export_service

    if _export_service is None:
        _export_service = AmendmentExportService()

    return _export_service
