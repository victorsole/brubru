"""
Document Exporter

Render user_documents (markdown content) as DOCX or PDF bytes.

DOCX path: parses markdown headings/lists/paragraphs into a python-docx Document.
PDF path: same parsed structure rendered with reportlab Platypus.

The renderer handles headings, paragraphs, bullet/numbered lists, bold/italic,
code blocks, blockquotes, markdown tables and images (embedded inline).
"""

from __future__ import annotations

import io
import re
from dataclasses import dataclass
from pathlib import Path
from typing import List, Optional


# ---------------------------------------------------------------------------
# Minimal markdown -> block parser
# ---------------------------------------------------------------------------

@dataclass
class Block:
    kind: str          # 'h1'|'h2'|'h3'|'p'|'ul'|'ol'|'code'|'quote'|'hr'|'table'|'image'
    text: str = ""
    items: Optional[List[str]] = None
    rows: Optional[List[List[str]]] = None  # table rows (header first)


_FENCE_RE = re.compile(r"^```")
_HEADING_RE = re.compile(r"^(#{1,6})\s+(.+?)\s*#*\s*$")
_UL_RE = re.compile(r"^\s*[-*]\s+(.+)$")
_OL_RE = re.compile(r"^\s*\d+[.)]\s+(.+)$")
_QUOTE_RE = re.compile(r"^>\s?(.*)$")
_HR_RE = re.compile(r"^\s*(-{3,}|_{3,}|\*{3,})\s*$")
_IMG_LINE_RE = re.compile(r"^!\[([^\]]*)\]\(([^)]+)\)$")


def _is_table_sep(line: str) -> bool:
    """A markdown table separator row, e.g. |---|:--:|---|."""
    s = line.strip()
    if "|" not in s or "-" not in s:
        return False
    return bool(re.fullmatch(r"[\s|:\-]+", s))


def _split_table_row(line: str) -> List[str]:
    s = line.strip().strip("|")
    return [c.strip() for c in s.split("|")]


def _load_image_bytes(url: str) -> Optional[bytes]:
    """Resolve an image URL to bytes for embedding. Returns None on any failure."""
    try:
        if not url:
            return None
        clean = url.split("?")[0]
        # Editor asset: read from local storage, no HTTP round-trip.
        if "/api/documents/assets/" in clean:
            asset_id = clean.rstrip("/").split("/api/documents/assets/")[-1]
            from services.storage.document_storage import get_document_storage
            return get_document_storage().get_document_content(asset_id)
        # Frontend static asset (e.g. /clients/bo.png).
        if "/clients/" in clean:
            rel = clean.split("/clients/")[-1]
            p = Path(__file__).resolve().parents[3] / "frontend" / "public" / "clients" / rel
            return p.read_bytes() if p.exists() else None
        # Remote image.
        if clean.startswith("http://") or clean.startswith("https://"):
            import httpx
            r = httpx.get(clean, timeout=10.0, follow_redirects=True)
            if r.status_code == 200 and len(r.content) <= 8 * 1024 * 1024:
                return r.content
        return None
    except Exception:
        return None


def _strip_fence_wrap(text: str) -> str:
    """Drop leading ```markdown / ``` fences if the whole content is wrapped."""
    stripped = text.strip()
    if stripped.startswith("```"):
        # Remove first line and any trailing fence
        lines = stripped.split("\n")
        if lines and lines[0].startswith("```"):
            lines = lines[1:]
        if lines and lines[-1].strip().startswith("```"):
            lines = lines[:-1]
        return "\n".join(lines)
    return text


def parse_markdown(text: str) -> List[Block]:
    """Best-effort markdown parser that keeps the structure we render."""
    if not text:
        return []

    text = _strip_fence_wrap(text)
    lines = text.replace("\r\n", "\n").split("\n")

    blocks: List[Block] = []
    i = 0
    n = len(lines)

    while i < n:
        line = lines[i]
        stripped = line.strip()

        # Blank line
        if not stripped:
            i += 1
            continue

        # Code fence
        if _FENCE_RE.match(stripped):
            i += 1
            buf: List[str] = []
            while i < n and not _FENCE_RE.match(lines[i].strip()):
                buf.append(lines[i])
                i += 1
            if i < n:
                i += 1  # consume closing fence
            blocks.append(Block("code", "\n".join(buf)))
            continue

        # Horizontal rule
        if _HR_RE.match(stripped):
            blocks.append(Block("hr"))
            i += 1
            continue

        # Standalone image line: ![alt](url)
        m_img = _IMG_LINE_RE.match(stripped)
        if m_img:
            blocks.append(Block("image", text=m_img.group(2).strip(), items=[m_img.group(1)]))
            i += 1
            continue

        # Table: a pipe row followed by a separator row
        if stripped.startswith("|") and i + 1 < n and _is_table_sep(lines[i + 1]):
            rows = [_split_table_row(stripped)]
            i += 2  # consume header + separator
            while i < n and lines[i].strip().startswith("|"):
                rows.append(_split_table_row(lines[i]))
                i += 1
            blocks.append(Block("table", rows=rows))
            continue

        # Heading
        m = _HEADING_RE.match(stripped)
        if m:
            level = min(len(m.group(1)), 3)
            blocks.append(Block(f"h{level}", m.group(2).strip()))
            i += 1
            continue

        # Bullet list
        if _UL_RE.match(line):
            items: List[str] = []
            while i < n and _UL_RE.match(lines[i]):
                items.append(_UL_RE.match(lines[i]).group(1).strip())
                i += 1
            blocks.append(Block("ul", items=items))
            continue

        # Numbered list
        if _OL_RE.match(line):
            items = []
            while i < n and _OL_RE.match(lines[i]):
                items.append(_OL_RE.match(lines[i]).group(1).strip())
                i += 1
            blocks.append(Block("ol", items=items))
            continue

        # Blockquote
        if _QUOTE_RE.match(stripped):
            buf = []
            while i < n and _QUOTE_RE.match(lines[i].strip()):
                buf.append(_QUOTE_RE.match(lines[i].strip()).group(1))
                i += 1
            blocks.append(Block("quote", " ".join(buf).strip()))
            continue

        # Paragraph (consume consecutive non-empty, non-special lines)
        buf = [stripped]
        i += 1
        while i < n:
            nxt = lines[i]
            nxt_strip = nxt.strip()
            if (
                not nxt_strip
                or _HEADING_RE.match(nxt_strip)
                or _UL_RE.match(nxt)
                or _OL_RE.match(nxt)
                or _QUOTE_RE.match(nxt_strip)
                or _FENCE_RE.match(nxt_strip)
                or _HR_RE.match(nxt_strip)
                or _IMG_LINE_RE.match(nxt_strip)
                or nxt_strip.startswith("|")
            ):
                break
            buf.append(nxt_strip)
            i += 1
        blocks.append(Block("p", " ".join(buf)))

    return blocks


# ---------------------------------------------------------------------------
# Inline markdown (bold / italic / code / links) helpers
# ---------------------------------------------------------------------------

_INLINE_RE = re.compile(
    r"(\*\*([^*]+)\*\*)"        # bold
    r"|(__([^_]+)__)"            # bold
    r"|(\*([^*]+)\*)"            # italic
    r"|(_([^_]+)_)"              # italic
    r"|(`([^`]+)`)"              # code
    r"|(\[([^\]]+)\]\(([^)]+)\))"  # link
)


def _iter_inline(text: str):
    """Yield (kind, text, href) tuples. kind: 'plain'|'bold'|'italic'|'code'|'link'."""
    pos = 0
    for m in _INLINE_RE.finditer(text):
        if m.start() > pos:
            yield ("plain", text[pos : m.start()], None)
        if m.group(1):
            yield ("bold", m.group(2), None)
        elif m.group(3):
            yield ("bold", m.group(4), None)
        elif m.group(5):
            yield ("italic", m.group(6), None)
        elif m.group(7):
            yield ("italic", m.group(8), None)
        elif m.group(9):
            yield ("code", m.group(10), None)
        elif m.group(11):
            yield ("link", m.group(12), m.group(13))
        pos = m.end()
    if pos < len(text):
        yield ("plain", text[pos:], None)


# ---------------------------------------------------------------------------
# DOCX renderer (python-docx)
# ---------------------------------------------------------------------------

def export_docx(title: str, markdown_text: str) -> bytes:
    """Render markdown to a DOCX file and return the bytes."""
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # Document title
    if title:
        heading = doc.add_heading(title, level=0)
        for run in heading.runs:
            run.font.size = Pt(20)

    def add_runs(paragraph, text: str) -> None:
        for kind, chunk, _href in _iter_inline(text):
            if not chunk:
                continue
            run = paragraph.add_run(chunk)
            if kind == "bold":
                run.bold = True
            elif kind == "italic":
                run.italic = True
            elif kind == "code":
                run.font.name = "Courier New"
            elif kind == "link":
                run.font.color.rgb = RGBColor(0x06, 0x93, 0xE3)
                run.font.underline = True

    for block in parse_markdown(markdown_text or ""):
        if block.kind == "h1":
            doc.add_heading(block.text, level=1)
        elif block.kind == "h2":
            doc.add_heading(block.text, level=2)
        elif block.kind == "h3":
            doc.add_heading(block.text, level=3)
        elif block.kind == "p":
            para = doc.add_paragraph()
            add_runs(para, block.text)
        elif block.kind == "ul":
            for item in block.items or []:
                p = doc.add_paragraph(style="List Bullet")
                add_runs(p, item)
        elif block.kind == "ol":
            for item in block.items or []:
                p = doc.add_paragraph(style="List Number")
                add_runs(p, item)
        elif block.kind == "code":
            p = doc.add_paragraph(block.text)
            for run in p.runs:
                run.font.name = "Courier New"
                run.font.size = Pt(10)
        elif block.kind == "quote":
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Pt(18)
            run = p.add_run(block.text)
            run.italic = True
        elif block.kind == "hr":
            p = doc.add_paragraph("_" * 40)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif block.kind == "table":
            rows = block.rows or []
            if rows:
                cols = max(len(r) for r in rows)
                table = doc.add_table(rows=len(rows), cols=cols)
                table.style = "Table Grid"
                for ri, row in enumerate(rows):
                    for ci in range(cols):
                        cell = table.cell(ri, ci)
                        cell.text = ""
                        cp = cell.paragraphs[0]
                        add_runs(cp, row[ci] if ci < len(row) else "")
                        if ri == 0:
                            for run in cp.runs:
                                run.bold = True
                doc.add_paragraph()
        elif block.kind == "image":
            data = _load_image_bytes(block.text)
            if data:
                from docx.shared import Inches
                try:
                    doc.add_picture(io.BytesIO(data), width=Inches(6.0))
                except Exception:
                    alt = (block.items or [""])[0]
                    doc.add_paragraph(f"[image: {alt}]" if alt else "[image]")

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


# ---------------------------------------------------------------------------
# PDF renderer (reportlab)
# ---------------------------------------------------------------------------

def _inline_to_rl(text: str) -> str:
    """Convert markdown inline tags to ReportLab's paraparser tags."""
    parts: List[str] = []
    for kind, chunk, href in _iter_inline(text):
        if not chunk:
            continue
        chunk = chunk.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
        if kind == "bold":
            parts.append(f"<b>{chunk}</b>")
        elif kind == "italic":
            parts.append(f"<i>{chunk}</i>")
        elif kind == "code":
            parts.append(f'<font face="Courier">{chunk}</font>')
        elif kind == "link":
            parts.append(f'<a href="{href}" color="#0693e3">{chunk}</a>')
        else:
            parts.append(chunk)
    return "".join(parts)


def export_pdf(title: str, markdown_text: str) -> bytes:
    """Render markdown to a PDF file and return the bytes."""
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.platypus import (
        SimpleDocTemplate,
        Paragraph,
        Spacer,
        ListFlowable,
        ListItem,
        Preformatted,
        HRFlowable,
        Table as RLTable,
        TableStyle,
        Image as RLImage,
    )
    from reportlab.lib.utils import ImageReader
    from reportlab.lib import colors

    buffer = io.BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=A4,
        leftMargin=2 * cm,
        rightMargin=2 * cm,
        topMargin=2 * cm,
        bottomMargin=2 * cm,
        title=title or "Document",
    )

    styles = getSampleStyleSheet()
    body = ParagraphStyle(
        "Body",
        parent=styles["BodyText"],
        fontName="Helvetica",
        fontSize=10.5,
        leading=14,
        spaceAfter=6,
    )
    h1 = ParagraphStyle(
        "H1", parent=styles["Heading1"], fontName="Helvetica-Bold", fontSize=18, spaceAfter=10, textColor=colors.HexColor("#1f2937")
    )
    h2 = ParagraphStyle(
        "H2", parent=styles["Heading2"], fontName="Helvetica-Bold", fontSize=14, spaceAfter=8, textColor=colors.HexColor("#1f2937")
    )
    h3 = ParagraphStyle(
        "H3", parent=styles["Heading3"], fontName="Helvetica-Bold", fontSize=12, spaceAfter=6, textColor=colors.HexColor("#1f2937")
    )
    quote = ParagraphStyle("Quote", parent=body, leftIndent=18, textColor=colors.HexColor("#4b5563"))

    story = []

    if title:
        title_style = ParagraphStyle(
            "Title",
            parent=styles["Title"],
            fontName="Helvetica-Bold",
            fontSize=22,
            spaceAfter=18,
            textColor=colors.HexColor("#0f172a"),
        )
        story.append(Paragraph(title.replace("&", "&amp;"), title_style))

    for block in parse_markdown(markdown_text or ""):
        if block.kind == "h1":
            story.append(Paragraph(_inline_to_rl(block.text), h1))
        elif block.kind == "h2":
            story.append(Paragraph(_inline_to_rl(block.text), h2))
        elif block.kind == "h3":
            story.append(Paragraph(_inline_to_rl(block.text), h3))
        elif block.kind == "p":
            story.append(Paragraph(_inline_to_rl(block.text), body))
        elif block.kind in ("ul", "ol"):
            items = [ListItem(Paragraph(_inline_to_rl(item), body)) for item in (block.items or [])]
            bullet_type = "bullet" if block.kind == "ul" else "1"
            story.append(ListFlowable(items, bulletType=bullet_type, leftIndent=18))
            story.append(Spacer(1, 4))
        elif block.kind == "code":
            story.append(Preformatted(block.text, styles["Code"]))
            story.append(Spacer(1, 4))
        elif block.kind == "quote":
            story.append(Paragraph(_inline_to_rl(block.text), quote))
        elif block.kind == "hr":
            story.append(HRFlowable(width="100%", color=colors.HexColor("#d1d5db"), spaceBefore=6, spaceAfter=6))
        elif block.kind == "table":
            rows = block.rows or []
            if rows:
                cols = max(len(r) for r in rows)
                data = []
                for ri, row in enumerate(rows):
                    cells = []
                    for ci in range(cols):
                        raw = row[ci] if ci < len(row) else ""
                        inner = _inline_to_rl(raw)
                        if ri == 0:
                            inner = f"<b>{inner}</b>"
                        cells.append(Paragraph(inner, body))
                    data.append(cells)
                tbl = RLTable(data, repeatRows=1)
                tbl.setStyle(TableStyle([
                    ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#cbd5e1")),
                    ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#f1f5f9")),
                    ("VALIGN", (0, 0), (-1, -1), "TOP"),
                    ("LEFTPADDING", (0, 0), (-1, -1), 5),
                    ("RIGHTPADDING", (0, 0), (-1, -1), 5),
                    ("TOPPADDING", (0, 0), (-1, -1), 3),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
                ]))
                story.append(tbl)
                story.append(Spacer(1, 6))
        elif block.kind == "image":
            data = _load_image_bytes(block.text)
            if data:
                try:
                    iw, ih = ImageReader(io.BytesIO(data)).getSize()
                    iw, ih = float(iw), float(ih)
                    if iw > 0 and ih > 0:
                        # Usable area, minus a buffer for the frame's internal padding
                        # (reportlab frames reserve ~6pt each side, so margin math alone overflows).
                        max_w = A4[0] - 4 * cm - 12
                        max_h = A4[1] - 4 * cm - 40
                        # Scale to fit both dimensions; never upscale.
                        scale = min(max_w / iw, max_h / ih, 1.0)
                        story.append(RLImage(io.BytesIO(data), width=iw * scale, height=ih * scale))
                        story.append(Spacer(1, 6))
                except Exception:
                    pass

    doc.build(story)
    return buffer.getvalue()
